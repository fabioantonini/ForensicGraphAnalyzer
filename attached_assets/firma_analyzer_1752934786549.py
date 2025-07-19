# === IMPORT LIBRERIE ===
import cv2
import numpy as np
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from PIL import Image, ImageTk
from skimage.metrics import structural_similarity as ssim
import matplotlib.pyplot as plt
import os
import math
from docx import Document

# === VARIABILI GLOBALI ===
DPI = 300

firme_verifica = []
firme_comparative = []
firme_verifica_images = []
firme_comparative_images = []

anagrafica = {}

# === FUNZIONI DI BASE ===

def inserisci_anagrafica():
    global anagrafica
    anagrafica['Nome Caso'] = simpledialog.askstring("Anagrafica", "Inserisci il nome del caso:")
    anagrafica['Periziando / Controparte'] = simpledialog.askstring("Anagrafica", "Inserisci il nome del periziando o controparte:")
    anagrafica['Data'] = simpledialog.askstring("Anagrafica", "Inserisci la data (es. 26/04/2025):")
    anagrafica['Tipo Documento'] = simpledialog.askstring("Anagrafica", "Inserisci il tipo di documento (es. Testamento, Contratto, Delega):")
    anagrafica['Note'] = simpledialog.askstring("Anagrafica", "Eventuali note aggiuntive:")

def load_signature(is_verifica):
    file_paths = filedialog.askopenfilenames(
        title="Seleziona firme",
        filetypes=[("Image Files", "*.png;*.jpg;*.jpeg;*.bmp;*.tiff")]
    )
    if not file_paths:
        return
    
    for file_path in file_paths:
        image = cv2.imread(file_path, cv2.IMREAD_GRAYSCALE)
        img_display = Image.open(file_path).resize((150, 75))
        img_display = ImageTk.PhotoImage(img_display)

        if is_verifica:
            firme_verifica.append((image, file_path))
            firme_verifica_images.append(img_display)
        else:
            firme_comparative.append((image, file_path))
            firme_comparative_images.append(img_display)
    
    update_canvas()

def update_canvas():
    canvas_verifica.delete("all")
    canvas_comparative.delete("all")

    for i, img in enumerate(firme_verifica_images):
        canvas_verifica.create_image((i * 160, 0), anchor="nw", image=img)
    
    for i, img in enumerate(firme_comparative_images):
        canvas_comparative.create_image((i * 160, 0), anchor="nw", image=img)

def pixels_to_mm(pixels):
    return (pixels * 25.4) / DPI

def preprocess_image(image):
    image = cv2.resize(image, (300, 150))
    _, thresh = cv2.threshold(image, 150, 255, cv2.THRESH_BINARY_INV)
    return thresh

def calculate_curvature(contour):
    angles = []
    for i in range(2, len(contour)):
        pt1 = contour[i - 2][0]
        pt2 = contour[i - 1][0]
        pt3 = contour[i][0]
        v1 = pt1 - pt2
        v2 = pt3 - pt2
        angle = math.degrees(math.acos(
            np.clip(np.dot(v1, v2) / (np.linalg.norm(v1) * np.linalg.norm(v2) + 1e-5), -1.0, 1.0)
        ))
        angles.append(angle)
    return np.mean(angles) if angles else 0

def calculate_circularity(cnt):
    area = cv2.contourArea(cnt)
    perimeter = cv2.arcLength(cnt, True)
    if perimeter == 0:
        return 0
    return 4 * math.pi * area / (perimeter ** 2)

def analyze_signature(image):
    contours, _ = cv2.findContours(image, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if not contours:
        return None

    x_min = min([cv2.boundingRect(cnt)[0] for cnt in contours])
    y_min = min([cv2.boundingRect(cnt)[1] for cnt in contours])
    x_max = max([cv2.boundingRect(cnt)[0] + cv2.boundingRect(cnt)[2] for cnt in contours])
    y_max = max([cv2.boundingRect(cnt)[1] + cv2.boundingRect(cnt)[3] for cnt in contours])
    w = x_max - x_min
    h = y_max - y_min
    dimensions = (pixels_to_mm(w), pixels_to_mm(h))
    
    proportion = w / h if h > 0 else 0
    inclination = cv2.fitEllipse(contours[0])[2] if len(contours[0]) >= 5 else 0
    pressure_mean = np.mean(image)
    pressure_std = np.std(image)
    curvature = np.mean([calculate_curvature(cnt) for cnt in contours])
    readability = "Alta" if pressure_mean > 90 else "Media" if pressure_mean > 60 else "Bassa"
    style = "Corsivo" if proportion > 2 else "Stampatello" if proportion < 1.2 else "Misto"

    internal_contours, _ = cv2.findContours(image, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
    asole = [cnt for cnt in internal_contours if 20 < cv2.contourArea(cnt) < 500 and calculate_circularity(cnt) > 0.5]
    avg_asola_size = np.mean([cv2.contourArea(a) for a in asole]) if asole else 0

    x_positions = [cv2.boundingRect(cnt)[0] for cnt in contours]
    x_positions.sort()
    spacings = [x_positions[i+1] - x_positions[i] for i in range(len(x_positions)-1)] if len(x_positions) > 1 else [0]
    avg_spacing = np.mean(spacings)

    total_length = sum([cv2.arcLength(cnt, False) for cnt in contours])
    straight_distance = math.hypot(x_max - x_min, y_max - y_min)
    velocity = total_length / (straight_distance + 1e-5)

    overlap_ratio = np.sum(image > 0) / (w * h) if w * h > 0 else 0

    letter_connections = len(contours)
    baseline_y_positions = [pt[0][1] for cnt in contours for pt in cnt]
    baseline_std = np.std(baseline_y_positions) if baseline_y_positions else 0
    baseline_mm_std = pixels_to_mm(baseline_std)

    return {
        'Dimensions': dimensions,
        'Proportion': proportion,
        'Inclination': inclination,
        'PressureMean': pressure_mean,
        'PressureStd': pressure_std,
        'Curvature': curvature,
        'Readability': readability,
        'Style': style,
        'AvgAsolaSize': avg_asola_size,
        'AvgSpacing': avg_spacing,
        'Velocity': velocity,
        'OverlapRatio': overlap_ratio,
        'LetterConnections': letter_connections,
        'BaselineStd': baseline_mm_std
    }
# === BLOCCO: Creazione grafico compatibilità ===
def crea_grafico_confronto(dati_verifica, dati_comp):
    parametri_numerici = [
        'Proportion', 'Inclination', 'PressureMean', 'PressureStd',
        'Curvature', 'AvgAsolaSize', 'AvgSpacing', 'Velocity',
        'OverlapRatio', 'LetterConnections', 'BaselineStd'
    ]

    differenze = []
    etichette = []

    for parametro in parametri_numerici:
        valore_v = dati_verifica.get(parametro, 0)
        valore_c = dati_comp.get(parametro, 0)
        differenza = abs(valore_v - valore_c)
        differenze.append(differenza)
        etichette.append(parametro)

    max_diff = max(differenze) if differenze else 1
    compatibilita_percentuale = [(1 - (diff / max_diff)) * 100 for diff in differenze]

    plt.figure(figsize=(12, 6))
    plt.barh(etichette, compatibilita_percentuale)
    plt.xlabel('Compatibilità (%)')
    plt.title('Grafico Compatibilità Parametri Firma')
    plt.xlim(0, 100)
    plt.grid(axis='x')
    plt.tight_layout()
    plt.show()

# === BLOCCO: Report descrittivo automatico ===
def crea_report_descrittivo(dati_verifica, dati_comp):
    descrizione = ""

    if dati_verifica['Velocity'] > dati_comp['Velocity'] + 0.2:
        descrizione += "La firma in verifica presenta una maggiore velocità di esecuzione rispetto alla comparativa.\n"
    elif dati_verifica['Velocity'] < dati_comp['Velocity'] - 0.2:
        descrizione += "La firma in verifica presenta una minore velocità di esecuzione rispetto alla comparativa.\n"
    else:
        descrizione += "La velocità di esecuzione delle firme risulta compatibile.\n"

    if abs(dati_verifica['Proportion'] - dati_comp['Proportion']) < 0.2:
        descrizione += "Le proporzioni tra altezza e larghezza risultano simili.\n"
    else:
        descrizione += "Le proporzioni tra altezza e larghezza mostrano differenze significative.\n"

    if abs(dati_verifica['PressureMean'] - dati_comp['PressureMean']) < 10:
        descrizione += "La pressione esercitata durante la firma è compatibile tra i due esemplari.\n"
    else:
        descrizione += "La pressione presenta differenze evidenti tra le firme.\n"

    if abs(dati_verifica['Inclination'] - dati_comp['Inclination']) < 5:
        descrizione += "L'inclinazione dei tratti risulta simile.\n"
    else:
        descrizione += "L'inclinazione dei tratti evidenzia differenze stilistiche.\n"

    if abs(dati_verifica['Curvature'] - dati_comp['Curvature']) < 15:
        descrizione += "La curvilineità/angolosità delle firme è coerente.\n"
    else:
        descrizione += "La curvilineità/angolosità differisce sensibilmente tra le firme.\n"

    if abs(dati_verifica['AvgSpacing'] - dati_comp['AvgSpacing']) < 5:
        descrizione += "La spaziatura tra le lettere appare omogenea.\n"
    else:
        descrizione += "La spaziatura tra le lettere evidenzia disomogeneità.\n"

    return descrizione

# === BLOCCO: Salvataggio Report Word ===
def salva_report_word():
    testo = text_result.get("1.0", tk.END)

    if not testo.strip():
        messagebox.showwarning("Nessun Report", "Non ci sono dati da salvare.")
        return

    file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                              filetypes=[("Word Documents", "*.docx")],
                                              title="Salva report come...")
    if file_path:
        doc = Document()
        doc.add_heading("FirmaAnalyzer Pro - Report Analisi Firme", 0)

        for line in testo.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

        doc.save(file_path)
        messagebox.showinfo("Salvataggio completato", f"Report salvato correttamente in:\n{file_path}")

# === BLOCCO: Confronto firme principali ===
def compare_signatures():
    if not firme_verifica or not firme_comparative:
        messagebox.showerror("Errore", "Caricare almeno una firma per verifica e una comparativa.")
        return

    report_text = f"ANAGRAFICA CASO:\n"
    for chiave, valore in anagrafica.items():
        report_text += f"{chiave}: {valore}\n"
    report_text += "\nRISULTATI CONFRONTO:\n\n"

    for img_verifica, path_verifica in firme_verifica:
        processed_verifica = preprocess_image(img_verifica)
        dati_verifica = analyze_signature(processed_verifica)

        for img_comp, path_comp in firme_comparative:
            processed_comp = preprocess_image(img_comp)
            dati_comp = analyze_signature(processed_comp)

            similarity, _ = ssim(processed_verifica, processed_comp, full=True)

            report_text += f"Verifica: {os.path.basename(path_verifica)}\n"
            report_text += f"Comparativa: {os.path.basename(path_comp)}\n"
            report_text += f"Somiglianza SSIM: {similarity*100:.2f}%\n"

            for chiave in dati_verifica.keys():
                valore_v = dati_verifica[chiave]
                valore_c = dati_comp[chiave]
                if isinstance(valore_v, (int, float)):
                    differenza = abs(valore_v - valore_c)
                    report_text += f"{chiave}: {valore_v:.2f} vs {valore_c:.2f} (Diff: {differenza:.2f})\n"
                else:
                    report_text += f"{chiave}: {valore_v} vs {valore_c}\n"

            descrizione = crea_report_descrittivo(dati_verifica, dati_comp)
            report_text += "\nDESCRIZIONE TECNICA:\n" + descrizione + "\n"
            crea_grafico_confronto(dati_verifica, dati_comp)

            report_text += "\n"

    text_result.delete("1.0", tk.END)
    text_result.insert(tk.END, report_text)

# === BLOCCO: Funzione rimuovere firme ===
def remove_last_signature(is_verifica):
    global firme_verifica, firme_comparative
    global firme_verifica_images, firme_comparative_images

    if is_verifica:
        if firme_verifica:
            firme_verifica.pop()
            firme_verifica_images.pop()
    else:
        if firme_comparative:
            firme_comparative.pop()
            firme_comparative_images.pop()
    
    update_canvas()
# === BLOCCO: Creazione finestra principale ===
root = tk.Tk()
root.title("FirmaAnalyzer Pro - Analisi Grafologica Forense")
root.geometry("1200x850")

# Frame per i pulsanti principali
frame = tk.Frame(root)
frame.pack(pady=10)

btn_anagrafica = tk.Button(frame, text="Compila Anagrafica", command=inserisci_anagrafica)
btn_anagrafica.grid(row=0, column=0, padx=10)

btn_load_verifica = tk.Button(frame, text="Carica Firme Verifica", command=lambda: load_signature(True))
btn_load_verifica.grid(row=0, column=1, padx=10)

btn_load_comparative = tk.Button(frame, text="Carica Firme Comparative", command=lambda: load_signature(False))
btn_load_comparative.grid(row=0, column=2, padx=10)

btn_compare = tk.Button(frame, text="Confronta Firme", command=compare_signatures)
btn_compare.grid(row=0, column=3, padx=10)

btn_salva_report = tk.Button(frame, text="Salva Report Word", command=salva_report_word)
btn_salva_report.grid(row=1, column=0, padx=10, pady=5)

btn_remove_verifica = tk.Button(frame, text="Rimuovi Ultima Verifica", command=lambda: remove_last_signature(True))
btn_remove_verifica.grid(row=1, column=1, padx=10, pady=5)

btn_remove_comparative = tk.Button(frame, text="Rimuovi Ultima Comparativa", command=lambda: remove_last_signature(False))
btn_remove_comparative.grid(row=1, column=2, padx=10, pady=5)

# === BLOCCO: Canvas per firme ===
tk.Label(root, text="Firme in Verifica:").pack()
canvas_verifica = tk.Canvas(root, width=1000, height=120, bg="white")
canvas_verifica.pack()

tk.Label(root, text="Firme Comparative:").pack()
canvas_comparative = tk.Canvas(root, width=1000, height=120, bg="white")
canvas_comparative.pack()

# === BLOCCO: Area risultati ===
tk.Label(root, text="Risultati Analisi:").pack()
text_result = tk.Text(root, height=20, width=140)
text_result.pack(pady=10)

# === BLOCCO: Avvio programma ===
root.mainloop()

